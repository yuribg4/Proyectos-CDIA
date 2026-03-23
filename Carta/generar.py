#!/usr/bin/env python3
import pandas as pd, matplotlib, matplotlib.pyplot as plt, io, os
matplotlib.use('Agg')
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Configuración ──────────────────────────────────────────
EXCEL_PATH  = "calificaciones_alumnos.xlsx"
OUTPUT_DIR  = "cartas_calificaciones"
ESCUELA     = "Conalep 169"
CICLO       = "2025–2026"
UMBRAL      = 6.0
MESES       = dict(zip(
    ["January","February","March","April","May","June",
     "July","August","September","October","November","December"],
    ["enero","febrero","marzo","abril","mayo","junio",
     "julio","agosto","septiembre","octubre","noviembre","diciembre"]))
FECHA       = datetime.now().strftime("%d de %B de %Y")
for en, es in MESES.items(): FECHA = FECHA.replace(en, es)

AZUL, ROJO, VERDE, GRIS = (RGBColor(0x1F,0x49,0x7D), RGBColor(0xC0,0x39,0x2B),
                            RGBColor(0x1E,0x7E,0x34), RGBColor(0x70,0x70,0x70))

# ── Helpers ────────────────────────────────────────────────
def bg(cell, color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),color)
    cell._tc.get_or_add_tcPr().append(shd)

def run(p, texto, bold=False, italic=False, size=11, color=None):
    r = p.add_run(texto)
    r.bold=bold; r.italic=italic; r.font.size=Pt(size); r.font.name='Arial'
    if color: r.font.color.rgb=color
    return r

def parrafo(doc, texto, bold=False, italic=False, size=11, color=None,
            align=WD_ALIGN_PARAGRAPH.LEFT, sp=6):
    p = doc.add_paragraph(); p.alignment=align; p.paragraph_format.space_after=Pt(sp)
    run(p, texto, bold, italic, size, color); return p

def celda_texto(cell, texto, bold=False, size=10, color=None,
                align=WD_ALIGN_PARAGRAPH.CENTER):
    p = cell.paragraphs[0]; p.alignment=align
    run(p, texto, bold=bold, size=size, color=color)

# ── Gráfica ────────────────────────────────────────────────
def grafica(alumno, materias, cals, promedio):
    colores = ['#1e7e34' if c>=9 else '#1F497D' if c>=7 else '#e6a817' if c>=6 else '#C0392B'
               for c in cals]
    fig, ax = plt.subplots(figsize=(7, 3.5))
    bars = ax.bar(materias, cals, color=colores, edgecolor='white', linewidth=0.8, zorder=3)
    ax.set_ylim(0, 10.5)
    ax.axhline(6,        color='#C0392B', linestyle='--', lw=1.2, alpha=0.7, label='Mínimo (6.0)', zorder=2)
    ax.axhline(promedio, color='#1F497D', lw=1.5, alpha=0.8, label=f'Promedio ({promedio:.1f})', zorder=2)
    ax.set_ylabel('Calificación', fontsize=10)
    ax.set_title(f'Calificaciones — {alumno}', fontsize=11, fontweight='bold', pad=10)
    ax.grid(axis='y', alpha=0.3, zorder=0); ax.set_facecolor('#f9f9f9')
    for b, v in zip(bars, cals):
        ax.text(b.get_x()+b.get_width()/2., b.get_height()+0.1, str(v),
                ha='center', va='bottom', fontsize=10, fontweight='bold')
    ax.legend(fontsize=8, loc='lower right'); plt.tight_layout()
    buf = io.BytesIO(); plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0); plt.close(); return buf

# ── Carta ──────────────────────────────────────────────────
def generar_carta(row, materias, output_dir):
    alumno   = row['Alumno']
    cals     = [row[m] for m in materias]
    promedio = sum(cals) / len(cals)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width=Cm(21.59); sec.page_height=Cm(27.94)
    for a in ('left_margin','right_margin','top_margin','bottom_margin'): setattr(sec,a,Cm(2.54))
    doc.styles['Normal'].font.name='Arial'; doc.styles['Normal'].font.size=Pt(11)

    # Encabezado azul
    t = doc.add_table(rows=1, cols=1); t.style='Table Grid'
    c = t.rows[0].cells[0]; bg(c,'1F497D'); c.paragraphs[0].clear()
    p = c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(2)
    run(p, ESCUELA, bold=True, size=14, color=RGBColor(0xFF,0xFF,0xFF))
    p2 = c.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after=Pt(6)
    run(p2, f"Dirección General  |  Primer Reporte de Calificaciones  |  Ciclo Escolar {CICLO}",
        size=9, color=RGBColor(0xCC,0xDD,0xEE))

    doc.add_paragraph()
    p_f = doc.add_paragraph(); p_f.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    p_f.paragraph_format.space_after=Pt(12)
    run(p_f, f"Tepic, Nayarit, a {FECHA}", size=10, color=GRIS)

    parrafo(doc, f"Padre/Madre de familia de {alumno}", bold=True, sp=2)
    parrafo(doc, "Presente", italic=True, color=GRIS, sp=14)
    parrafo(doc, f"Estimado(a) Padre/Madre de Familia del alumno(a) {alumno}:", sp=10)

    doc.add_paragraph(
        "Por medio del presente comunicado, nos dirigimos a usted con el más sincero respeto, "
        f"con la finalidad de informarle sobre el desempeño académico de su hijo(a) "
        f"correspondiente al primer reporte de calificaciones del ciclo escolar {CICLO}. "
        "En esta institución nos comprometemos con la formación integral de cada alumno y "
        "consideramos esencial mantener una comunicación constante y transparente con las familias."
    ).paragraph_format.space_after = Pt(8)
    doc.add_paragraph(
        "A continuación, se presentan las calificaciones obtenidas en cada una de las asignaturas:"
    ).paragraph_format.space_after = Pt(10)

    # Tabla de calificaciones
    tbl = doc.add_table(rows=1+len(materias)+1, cols=3); tbl.style='Table Grid'
    for i, (h, w) in enumerate(zip(['Asignatura','Calificación','Estado'],
                                   [Cm(7), Cm(3.5), Cm(3)])):
        c = tbl.rows[0].cells[i]; bg(c,'1F497D'); c.width=w
        celda_texto(c, h, bold=True, color=RGBColor(0xFF,0xFF,0xFF))

    for i, materia in enumerate(materias):
        cal = row[materia]; cells = tbl.rows[i+1].cells; fondo = 'F2F2F2' if i%2==0 else 'FFFFFF'
        for c in cells: bg(c, fondo)
        celda_texto(cells[0], materia, align=WD_ALIGN_PARAGRAPH.LEFT)
        celda_texto(cells[1], str(cal))
        celda_texto(cells[2], "Aprobado ✓" if cal>=6 else "Reprobado ✗",
                    color=VERDE if cal>=6 else ROJO)

    for c in tbl.rows[-1].cells: bg(c,'D9E1F2')
    celda_texto(tbl.rows[-1].cells[0], "PROMEDIO GENERAL", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
    celda_texto(tbl.rows[-1].cells[1], f"{promedio:.1f}", bold=True, size=11,
                color=VERDE if promedio>=6 else ROJO)
    celda_texto(tbl.rows[-1].cells[2],
                "Satisfactorio" if promedio>=7 else "Suficiente" if promedio>=6 else "En riesgo",
                bold=True, color=VERDE if promedio>=7 else (AZUL if promedio>=6 else ROJO))

    doc.add_paragraph()
    parrafo(doc, "Representación gráfica del desempeño académico:", bold=True, color=AZUL, sp=6)
    doc.add_picture(grafica(alumno, materias, cals, promedio), width=Inches(5.8))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Recomendaciones si bajo rendimiento
    if promedio < UMBRAL:
        t_a = doc.add_table(rows=1, cols=1); t_a.style='Table Grid'
        ca = t_a.rows[0].cells[0]; bg(ca,'FDECEA')
        celda_texto(ca, "⚠  ATENCIÓN: El alumno presenta promedio por debajo del mínimo aprobatorio  ⚠",
                    bold=True, color=ROJO)
        doc.add_paragraph()
        parrafo(doc, "Recomendaciones para mejorar el desempeño académico:", bold=True, color=ROJO, sp=6)
        recs = [
            "Establecer un horario de estudio diario de al menos 2 horas en casa.",
            "Solicitar asesoría adicional con los maestros de materias reprobadas.",
            "Fomentar hábitos de lectura y repaso de apuntes de manera constante.",
            "Agendar una reunión con el tutor escolar para diseñar un plan de mejora.",
        ] + [f"En {m} (cal: {row[m]}): práctica adicional y revisión del primer bimestre."
             for m in materias if row[m] < UMBRAL]
        for rec in recs:
            p_r = doc.add_paragraph(style='List Bullet'); p_r.paragraph_format.space_after=Pt(4)
            run(p_r, rec, size=10)
        doc.add_paragraph()

    doc.add_paragraph(
        "Hacemos un cordial llamado a fortalecer el acompañamiento en casa, motivando al alumno "
        "a mantener una actitud positiva frente al aprendizaje."
    ).paragraph_format.space_after = Pt(8)
    doc.add_paragraph(
        "Agradecemos su confianza y quedamos a sus órdenes para cualquier aclaración."
    ).paragraph_format.space_after = Pt(14)

    parrafo(doc, "Atentamente,", sp=2)
    parrafo(doc, "La Dirección Escolar", italic=True, color=GRIS, sp=30)

    # Firma
    tf = doc.add_table(rows=1, cols=1); tf.style='Table Grid'
    cf = tf.rows[0].cells[0]; bg(cf,'F5F5F5')
    celda_texto(cf, "_________________________________", size=10)
    p_s = cf.add_paragraph(); p_s.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run(p_s, "Firma y Sello Institucional", size=9, color=GRIS)

    # Pie con borde superior
    doc.add_paragraph()
    pie = doc.add_paragraph(); pie.alignment=WD_ALIGN_PARAGRAPH.CENTER
    pie.paragraph_format.space_before=Pt(10)
    pBdr = OxmlElement('w:pBdr'); top = OxmlElement('w:top')
    for k,v in [('w:val','single'),('w:sz','6'),('w:space','4'),('w:color','1F497D')]:
        top.set(qn(k),v)
    pBdr.append(top); pie._p.get_or_add_pPr().append(pBdr)
    run(pie, f"{ESCUELA}  |  Ciclo Escolar {CICLO}  |  Documento generado el {FECHA}",
        size=8, color=GRIS)

    safe = alumno.translate(str.maketrans(' áéíóúñ','_aeioun'))
    path = os.path.join(output_dir, f"Carta_{safe}.docx")
    doc.save(path); return path, promedio

# ── Main ───────────────────────────────────────────────────
def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    df = pd.read_excel(EXCEL_PATH)
    materias = [c for c in df.columns if c != 'Alumno']
    print(f"\n{'='*50}\n  {ESCUELA} — Generador de Cartas\n{'='*50}\n")
    resultados = []
    for _, row in df.iterrows():
        path, prom = generar_carta(row, materias, OUTPUT_DIR)
        print(f"  ✓  {row['Alumno']:<22}  Promedio: {prom:.1f}{'  ⚠' if prom<UMBRAL else ''}")
        resultados.append(prom)
    print(f"\n  {len(resultados)} cartas en ./{OUTPUT_DIR}/  |  "
          f"Bajo rendimiento: {sum(p<UMBRAL for p in resultados)}\n{'='*50}\n")

if __name__ == "__main__":
    main()
